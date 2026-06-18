# -*- coding: utf-8 -*-
"""
================================================================================
DOCX Export Service
================================================================================
Kreira Word dokument od prevedenih chunk-ova.

Verzija: 1.0.0
================================================================================
"""

import io
import logging
from typing import List, Dict, Any
from datetime import datetime

from app.services.base_export_service import BaseExportService

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    DOCX_AVAILABLE = True
except ImportError:
    Document = None
    Pt = None
    RGBColor = None
    WD_ALIGN_PARAGRAPH = None
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


class DOCXExportService(BaseExportService):
    """Service za export dokumenta u DOCX format."""

    def generate(
        self,
        title: str,
        chunks: List[Dict[str, Any]],
        include_original: bool = False,
        use_skill: bool = True,
    ) -> bytes:
        """
        Generiše DOCX dokument od chunk-ova.

        Args:
            title: Naslov dokumenta
            chunks: Lista chunk-ova (dictionaries)
            include_original: Da li uključuje originalni tekst
            use_skill: Da li koristiti skill instrukcije za formatiranje

        Returns:
            DOCX bytes
        """
        self._report_progress(0, 100, "Priprema DOCX generisanja...")

        skill_prompt = self._load_skill("docx") if use_skill else ""
        if skill_prompt:
            logger.info(f"Loaded DOCX skill prompt: {skill_prompt[:100]}...")
        skill_options = self._parse_skill_prompt(skill_prompt)

        doc = Document()

        # ── Title ────────────────────────────────────────────────────────────
        title_paragraph = doc.add_heading(title, level=0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Primeni font na naslov
        for run in title_paragraph.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(16)

        # ── Table of Contents (ako je omogućeno) ────────────────────────────
        if skill_options.get("table_of_contents"):
            doc.add_heading("SADRŽAJ", level=1)
            for i, chunk in enumerate(chunks[:30], 1):
                trans = chunk.get("translated_content") or chunk.get(
                    "translated_text", ""
                )
                first_line = trans[:60].replace("\n", " ") if trans else ""
                p = doc.add_paragraph(f"{i}. {first_line}...")
                p.paragraph_format.space_after = Pt(3)

        # ── Metadata ────────────────────────────────────────────────────────
        meta_para = doc.add_paragraph(f"Broj sekcija: {len(chunks)}")
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_para.runs[0].font.italic = True

        self._report_progress(10, 100, "Kreiranje sadržaja...")

        # ── Chapters/Sections ──────────────────────────────────────────────
        bold_headings = skill_options.get("bold_headings", True)

        for idx, chunk in enumerate(chunks):
            # Odredi nivo headinga
            heading_level = chunk.get("heading_level", 1)
            if heading_level > 3:
                heading_level = 3

            # Koristi parent_heading kao naslov
            if chunk.get("parent_heading"):
                doc.add_heading(chunk.get("parent_heading"), level=2)

            content = self._get_content(chunk)

            if content:
                # Dodaj kao paragraph
                para = doc.add_paragraph(content)

                # Font settings
                font_name = skill_options.get("font", "Calibri")
                font_size = skill_options.get("font_size", 11)
                for run in para.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    if bold_headings:
                        run.font.bold = True

                # Justified text
                if skill_options.get("justified"):
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Original text
            if include_original and chunk.get("content"):
                orig_para = doc.add_paragraph(f"Original: {chunk['content'][:200]}...")
                orig_para.runs[0].font.size = Pt(9)
                orig_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            # Progress update
            if idx % 500 == 0:
                progress = 10 + int((idx / len(chunks)) * 80)
                self._report_progress(progress, 100, f"Sadržaj: {idx}/{len(chunks)}...")

        # ── Footer ────────────────────────────────────────────────────────
        self._report_progress(90, 100, "Finalizacija DOCX-a...")
        doc.add_paragraph()
        footer = doc.add_paragraph("---")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(
            f"Generisano: AI Learning System | {datetime.now().strftime('%d.%m.%Y')}",
            style="Intense Quote",
        )

        # Sačuvaj u bytes
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        self._report_progress(100, 100, "DOCX generisanje završeno!")
        return docx_bytes.getvalue()

    def _parse_skill_prompt(self, prompt: str) -> dict:
        """Parsira skill prompt i ekstrahuje opcije za DOCX formatiranje."""
        options = {
            "table_of_contents": False,
            "bold_headings": True,
            "justified": True,
            "font": "Calibri",
            "font_size": 11,
        }

        if not prompt:
            return options

        prompt_lower = prompt.lower()

        if (
            "table of contents" in prompt_lower
            or "sadržaj" in prompt_lower
            or "toc" in prompt_lower
        ):
            options["table_of_contents"] = True

        if "bold" in prompt_lower:
            options["bold_headings"] = True

        if "justified" in prompt_lower:
            options["justified"] = True

        if "arial" in prompt_lower:
            options["font"] = "Arial"
        elif "times" in prompt_lower:
            options["font"] = "Times New Roman"

        return options


docx_export_service = DOCXExportService()
