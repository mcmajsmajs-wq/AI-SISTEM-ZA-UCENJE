# -*- coding: utf-8 -*-
"""
================================================================================
Petar II Petrović-Njegoš
"Blago tome ko dovijek živi, imao se rašta i roditi"
================================================================================

AI Learning System
File Processing Service
Verzija: 1.0.0
Autor: Branko Suznjevic
Datum: 2026

Funkcionalnosti:
- PDF processing
- TXT processing
- DOCX processing
- Image processing with OCR
================================================================================
"""

import io
import logging
from typing import Dict, Any

logger = logging.getLogger(__name__)

# Check availability of optional libraries
TESSERACT_AVAILABLE = True
try:
    import pytesseract
    from PIL import Image
except ImportError:
    TESSERACT_AVAILABLE = False
    logger.warning("Tesseract OCR not available - image OCR disabled")

DOCX_AVAILABLE = True
try:
    import docx
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available - DOCX processing disabled")


class FileProcessingService:
    """
    FILE PROCESSING SERVICE
    Processes various file types and extracts text content.
    """

    def __init__(self, use_ocr: bool = True, ocr_language: str = "srp"):
        self.use_ocr = use_ocr and TESSERACT_AVAILABLE
        self.ocr_language = ocr_language

    def process_file(
        self, file_content: bytes, filename: str, file_ext: str, progress_callback=None
    ) -> Dict[str, Any]:
        """
        Process a file and extract text based on its type.

        Args:
            file_content: Raw file bytes
            filename: Original filename
            file_ext: File extension (with dot)
            progress_callback: Callback for progress tracking

        Returns:
            Dict with 'success', 'text', 'metadata'
        """
        file_ext = file_ext.lower()

        try:
            if file_ext == ".pdf":
                return self._process_pdf(file_content, progress_callback)
            elif file_ext == ".txt":
                return self._process_txt(file_content)
            elif file_ext == ".docx":
                return self._process_docx(file_content)
            elif file_ext in {
                ".jpg",
                ".jpeg",
                ".png",
                ".gif",
                ".bmp",
                ".tiff",
                ".webp",
            }:
                return self._process_image(file_content, file_ext)
            else:
                return {
                    "success": False,
                    "text": "",
                    "error": f"Unsupported file type: {file_ext}",
                    "metadata": {},
                }
        except Exception as e:
            logger.error(f"Error processing file {filename}: {e}")
            return {"success": False, "text": "", "error": str(e), "metadata": {}}

    def _process_pdf(
        self, file_content: bytes, progress_callback=None
    ) -> Dict[str, Any]:
        from app.services.pdf import PDFService

        pdf_service = PDFService(use_ocr=self.use_ocr, ocr_language=self.ocr_language)
        result = pdf_service.process_pdf(
            file_content, progress_callback=progress_callback
        )

        if not result.success:
            return {"success": False, "text": "", "error": result.error, "metadata": {}}

        # Combine all chunks into single text
        all_text = "\n\n".join(chunk.content for chunk in result.chunks)

        return {
            "success": True,
            "text": all_text,
            "metadata": {
                "total_pages": result.metadata.total_pages,
                "total_chunks": len(result.chunks),
                "is_scanned": result.metadata.is_scanned,
                "has_images": result.metadata.has_images,
            },
        }

    def _process_txt(self, file_content: bytes) -> Dict[str, Any]:
        try:
            # Try UTF-8 first, then fallback to other encodings
            text = file_content.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text = file_content.decode("latin-1")
            except UnicodeDecodeError:
                text = file_content.decode("utf-8", errors="ignore")

        # Clean up the text
        text = text.strip()

        return {
            "success": True,
            "text": text,
            "metadata": {"char_count": len(text), "word_count": len(text.split())},
        }

    def _process_docx(self, file_content: bytes) -> Dict[str, Any]:
        if not DOCX_AVAILABLE:
            return {
                "success": False,
                "text": "",
                "error": "DOCX processing not available. Install python-docx.",
                "metadata": {},
            }

        try:
            doc = docx.Document(io.BytesIO(file_content))
            paragraphs = [
                para.text.strip() for para in doc.paragraphs if para.text.strip()
            ]
            text = "\n\n".join(paragraphs)

            # Also extract tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        tables_text.append(row_text)

            if tables_text:
                text += "\n\n--- Tables ---\n\n"
                text += "\n".join(tables_text)

            return {
                "success": True,
                "text": text,
                "metadata": {
                    "paragraphs": len(paragraphs),
                    "tables": len(doc.tables),
                    "char_count": len(text),
                },
            }
        except Exception as e:
            return {
                "success": False,
                "text": "",
                "error": f"Failed to process DOCX: {str(e)}",
                "metadata": {},
            }

    def _process_image(self, file_content: bytes, file_ext: str) -> Dict[str, Any]:
        if not self.use_ocr:
            return {
                "success": False,
                "text": "",
                "error": "OCR is disabled",
                "metadata": {},
            }

        try:
            image = Image.open(io.BytesIO(file_content))

            # Perform OCR
            text = pytesseract.image_to_string(image, lang=self.ocr_language)
            text = text.strip()

            # Get metadata
            width, height = image.size
            format_name = image.format

            return {
                "success": True,
                "text": text,
                "metadata": {
                    "width": width,
                    "height": height,
                    "format": format_name,
                    "mode": image.mode,
                    "char_count": len(text),
                },
            }
        except Exception as e:
            return {
                "success": False,
                "text": "",
                "error": f"OCR failed: {str(e)}",
                "metadata": {},
            }


# Singleton instance
file_processing_service = FileProcessingService()
